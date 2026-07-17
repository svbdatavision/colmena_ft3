"""
Utilidades del Agente IA para Clasificación de NDAs.

Este módulo concentra:
- Lectura de archivos PDF y DOCX.
- Carga automática de modelos desde la carpeta docs/.
- Inferencia de tipología a partir del nombre del archivo.
- Clasificación por similitud TF-IDF + cosine similarity.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import BinaryIO, Dict, List, Union

import numpy as np
import pdfplumber
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# Extensiones de documento admitidas por el clasificador.
ALLOWED_EXTENSIONS = {".pdf", ".docx"}

# Carpeta por defecto donde viven los modelos de referencia.
DEFAULT_DOCS_DIR = Path(__file__).resolve().parent / "docs"


def extract_text_from_pdf(source: Union[str, Path, BinaryIO, bytes]) -> str:
    """
    Extrae todo el texto de un archivo PDF utilizando pdfplumber.

    Parameters
    ----------
    source:
        Ruta al PDF, objeto tipo archivo (BytesIO/UploadedFile) o bytes.

    Returns
    -------
    str
        Texto completo concatenado de todas las páginas.
    """
    # Normalizamos la entrada a un buffer legible por pdfplumber.
    if isinstance(source, (str, Path)):
        pdf_file = open(source, "rb")
        should_close = True
    elif isinstance(source, bytes):
        pdf_file = io.BytesIO(source)
        should_close = False
    else:
        # Streamlit UploadedFile u otros file-like objects.
        if hasattr(source, "seek"):
            source.seek(0)
        pdf_file = source
        should_close = False

    pages_text: List[str] = []
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                # extract_text() puede devolver None en páginas vacías.
                page_text = page.extract_text() or ""
                pages_text.append(page_text)
    finally:
        if should_close:
            pdf_file.close()

    return "\n".join(pages_text).strip()


def extract_text_from_docx(source: Union[str, Path, BinaryIO, bytes]) -> str:
    """
    Extrae todo el texto de un archivo DOCX utilizando python-docx.

    Incluye párrafos del cuerpo y, cuando existen, celdas de tablas.

    Parameters
    ----------
    source:
        Ruta al DOCX, objeto tipo archivo o bytes.

    Returns
    -------
    str
        Texto completo del documento.
    """
    if isinstance(source, (str, Path)):
        document = Document(str(source))
    elif isinstance(source, bytes):
        document = Document(io.BytesIO(source))
    else:
        if hasattr(source, "seek"):
            source.seek(0)
        document = Document(source)

    chunks: List[str] = []

    # Párrafos del cuerpo principal.
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    # Texto contenido en tablas (frecuente en plantillas legales).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    chunks.append(cell_text)

    return "\n".join(chunks).strip()


def extract_text_from_file(
    source: Union[str, Path, BinaryIO, bytes],
    filename: str | None = None,
) -> str:
    """
    Detecta la extensión y delega la extracción a PDF o DOCX.

    Parameters
    ----------
    source:
        Contenido o ruta del documento a leer.
    filename:
        Nombre original del archivo (útil con Streamlit UploadedFile).

    Returns
    -------
    str
        Texto extraído.

    Raises
    ------
    ValueError
        Si la extensión no es PDF ni DOCX, o si el texto queda vacío.
    """
    # Determinamos la extensión a partir del nombre o de la ruta.
    if filename:
        extension = Path(filename).suffix.lower()
    elif isinstance(source, (str, Path)):
        extension = Path(source).suffix.lower()
    else:
        raise ValueError(
            "No se pudo determinar la extensión del archivo. "
            "Indique el nombre del archivo (filename)."
        )

    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Formato no soportado: '{extension}'. "
            "Solo se admiten archivos PDF y DOCX."
        )

    if extension == ".pdf":
        text = extract_text_from_pdf(source)
    else:
        text = extract_text_from_docx(source)

    if not text:
        raise ValueError(
            "No se pudo extraer texto del documento. "
            "Verifique que el archivo no esté vacío o escaneado solo como imagen."
        )

    return text


def infer_nda_type(model_name: str) -> str:
    """
    Infiere una tipología legible a partir del nombre del modelo.

    Se usan palabras clave frecuentes en los nombres de plantillas NDA.
    """
    name = model_name.lower()

    if "unilateral" in name:
        return "NDA Unilateral"
    if "bilateral" in name or "robusto" in name:
        return "NDA Bilateral"
    if "procurement" in name:
        return "NDA Simplificado – Procurement"
    if "simplificado" in name or "simplificat" in name:
        return "NDA Simplificado"
    if "webdox" in name:
        return "NDA Simplificado – Webdox"
    if "mutuo" in name or "mutual" in name:
        return "NDA Mutuo / Bilateral"

    return "Acuerdo de Confidencialidad (genérico)"


def load_model_documents(docs_dir: Union[str, Path] = DEFAULT_DOCS_DIR) -> Dict[str, str]:
    """
    Lee automáticamente todos los PDF/DOCX presentes en la carpeta de modelos.

    Parameters
    ----------
    docs_dir:
        Ruta a la carpeta docs/ con los modelos de referencia.

    Returns
    -------
    dict
        Mapa {nombre_archivo: texto_extraído}.

    Raises
    ------
    FileNotFoundError
        Si la carpeta no existe.
    ValueError
        Si no hay documentos válidos o ninguno tiene texto usable.
    """
    docs_path = Path(docs_dir)
    if not docs_path.exists() or not docs_path.is_dir():
        raise FileNotFoundError(
            f"No se encontró la carpeta de modelos: '{docs_path}'. "
            "Cree la carpeta docs/ y coloque allí los modelos NDA."
        )

    models: Dict[str, str] = {}

    # Ordenamos para que el ranking sea determinista entre ejecuciones.
    candidate_files = sorted(
        [
            path
            for path in docs_path.iterdir()
            if path.is_file() and path.suffix.lower() in ALLOWED_EXTENSIONS
        ],
        key=lambda p: p.name.lower(),
    )

    if not candidate_files:
        raise ValueError(
            f"La carpeta '{docs_path}' no contiene archivos PDF o DOCX."
        )

    errors: List[str] = []
    for file_path in candidate_files:
        try:
            text = extract_text_from_file(file_path)
            models[file_path.name] = text
        except Exception as exc:  # noqa: BLE001 - reportamos y seguimos
            errors.append(f"{file_path.name}: {exc}")

    if not models:
        detail = "; ".join(errors) if errors else "sin detalle"
        raise ValueError(
            "No se pudo leer ningún modelo de la carpeta docs/. "
            f"Detalle: {detail}"
        )

    return models


def classify_document(
    uploaded_text: str,
    model_texts: Dict[str, str],
    top_n: int = 6,
) -> Dict:
    """
    Clasifica un NDA por similitud TF-IDF + cosine similarity.

    Flujo:
    1. Vectoriza el documento cargado junto con todos los modelos.
    2. Calcula cosine_similarity entre el documento y cada modelo.
    3. Selecciona el modelo más parecido.
    4. Devuelve ranking, confianza y tipología inferida.

    Parameters
    ----------
    uploaded_text:
        Texto del documento subido por el usuario.
    model_texts:
        Diccionario {nombre_modelo: texto}.
    top_n:
        Cantidad máxima de modelos a incluir en el ranking.

    Returns
    -------
    dict
        Resultado con modelo recomendado, similitudes y justificación.
    """
    if not uploaded_text or not uploaded_text.strip():
        raise ValueError("El texto del documento cargado está vacío.")

    if not model_texts:
        raise ValueError("No hay modelos de referencia para comparar.")

    model_names = list(model_texts.keys())
    corpus = [uploaded_text] + [model_texts[name] for name in model_names]

    # TF-IDF en español: quitamos stopwords comunes vía tokenización básica.
    # Usamos unigramas y bigramas para capturar cláusulas frecuentes.
    vectorizer = TfidfVectorizer(
        lowercase=True,
        analyzer="word",
        ngram_range=(1, 2),
        min_df=1,
        max_df=0.95,
        # Tokenizador simple: letras y números (útil con textos legales).
        token_pattern=r"(?u)\b\w+\b",
    )

    tfidf_matrix = vectorizer.fit_transform(corpus)

    # Fila 0 = documento cargado; filas 1..N = modelos.
    uploaded_vector = tfidf_matrix[0:1]
    models_matrix = tfidf_matrix[1:]

    similarities = cosine_similarity(uploaded_vector, models_matrix).flatten()

    # Ranking completo ordenado de mayor a menor similitud.
    ranked_indices = np.argsort(similarities)[::-1]
    ranking: List[Dict[str, Union[str, float]]] = []
    for index in ranked_indices:
        score = float(similarities[index])
        ranking.append(
            {
                "modelo": model_names[index],
                "similitud": score,
                "similitud_pct": round(score * 100.0, 2),
                "tipo": infer_nda_type(model_names[index]),
            }
        )

    best = ranking[0]
    confidence_pct = float(best["similitud_pct"])

    # Nivel de confianza cualitativo para la tarjeta de resultados.
    if confidence_pct >= 70:
        confidence_level = "Alta"
    elif confidence_pct >= 40:
        confidence_level = "Media"
    else:
        confidence_level = "Baja"

    # Justificación fija del MVP (demostración de concepto).
    justification = (
        "El documento presenta una alta similitud con el modelo seleccionado, "
        "considerando el vocabulario y las cláusulas presentes."
    )

    return {
        "modelo_recomendado": best["modelo"],
        "similitud": float(best["similitud"]),
        "similitud_pct": confidence_pct,
        "nivel_confianza": confidence_level,
        "tipo_detectado": best["tipo"],
        "justificacion": justification,
        "ranking": ranking[:top_n],
        "ranking_completo": ranking,
    }


def sanitize_label(name: str, max_length: int = 48) -> str:
    """
    Acorta nombres de archivo largos para etiquetas de gráficos.
    """
    clean = re.sub(r"\.(docx|pdf)$", "", name, flags=re.IGNORECASE)
    if len(clean) <= max_length:
        return clean
    return clean[: max_length - 1] + "…"


def build_similarity_dataframe(ranking: List[Dict]) -> "object":
    """
    Construye un DataFrame de pandas listo para visualizar en Streamlit.

    Se importa pandas aquí para mantener utils usable también sin Streamlit.
    """
    import pandas as pd

    rows = []
    for item in ranking:
        rows.append(
            {
                "Modelo": sanitize_label(str(item["modelo"])),
                "Modelo completo": item["modelo"],
                "Similitud (%)": item["similitud_pct"],
                "Tipo": item["tipo"],
            }
        )

    return pd.DataFrame(rows)
